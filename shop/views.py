from django.shortcuts import render, redirect, get_object_or_404
from .models import User, Product, Category, Cart, Order, Role, PaymentMethod, Warehouse, Review, WhatInOrder, ProductsOnWarehouse, OrderStatus
from django.http import JsonResponse
from .forms import UserForm, ProductForm, CategoryForm, OrderForm
from .forms import RegistrationForm
from functools import wraps
from django.utils.timezone import now
from django.db.models import F
from .models import SupportMessage
from .forms import SupportMessageForm
from django.conf import settings
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from .forms import ReviewForm
from django.contrib.admin.views.decorators import staff_member_required
from django.core.management import call_command
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
import datetime
import os
from django.db.models import Q
from django.db import models 
from django.core.exceptions import FieldError 
from django.forms.models import modelform_factory
from django.contrib.auth.hashers import make_password
from django.contrib.auth.hashers import check_password
from django.contrib.auth import login
from django.core.mail import send_mail
from django.template.loader import render_to_string
import qrcode
from io import BytesIO
import requests
from .decorators import role_required
import matplotlib.pyplot as plt
import base64
from django.db.models import Sum
from docx import Document
import logging
import csv

logger = logging.getLogger(__name__)
MODEL_MAP = {
    'users': User,
    'products': Product,
    'categories': Category,
    'orders': Order,
    'roles': Role,
    'payment_methods': PaymentMethod,
    'warehouses': Warehouse,
    'reviews': Review,
    'products_on_warehouse': ProductsOnWarehouse,
    'support_messages': SupportMessage,
    'what_in_order': WhatInOrder,
}

def role_required(role_name):
    def decorator(view_func):
        @wraps(view_func)
        def _wrapped_view(request, *args, **kwargs):
            user_id = request.session.get('user_id')
            if not user_id:
                return redirect('login')
            user = User.objects.get(id=user_id)
            if user.role.role_name != role_name:
                return redirect('dashboard') 
            return view_func(request, *args, **kwargs)
        return _wrapped_view
    return decorator

@role_required('admin')
def save_database(request):
    timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S')
    backup_file = os.path.join(settings.BASE_DIR, f"backup_{timestamp}.json")
    
    os.system(f'python manage.py dumpdata > "{backup_file}"')
    
    with open(backup_file, 'rb') as f:
        response = HttpResponse(f.read(), content_type='application/json')
        response['Content-Disposition'] = f'attachment; filename=backup_{timestamp}.json'
    os.remove(backup_file) 
    return response

@role_required('admin')
@csrf_exempt  
def load_database(request):
    if request.method == 'POST':
       
        backup_file = request.FILES.get('backup_file')
        
        if not backup_file:
            return JsonResponse({'error': 'Файл не загружен'}, status=400)
        
      
        temp_file_path = os.path.join(settings.MEDIA_ROOT, 'temp_backup.json')
        with open(temp_file_path, 'wb') as temp_file:
            for chunk in backup_file.chunks():
                temp_file.write(chunk)
        
        
        try:
            call_command('loaddata', temp_file_path)
            os.remove(temp_file_path)
            return JsonResponse({'success': 'Данные успешно загружены!'})
        except Exception as e:
            os.remove(temp_file_path)  
            return JsonResponse({'error': f'Ошибка загрузки данных: {str(e)}'}, status=500)
    
    return render(request, 'shop/load_database.html')


def register(request):
    if request.method == 'POST':
        form = RegistrationForm(request.POST)
        if form.is_valid():
            user = form.save(commit=False)
            user.password = make_password(form.cleaned_data['password'])  
            
            buyer_role, created = Role.objects.get_or_create(role_name='buyer')
            user.role = buyer_role
            
            user.save()
            return redirect('login')
    else:
        form = RegistrationForm()
    
    return render(request, 'shop/register.html', {'form': form})


def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = User.objects.filter(user_name=username).first()
        if user and check_password(password, user.password): 
            request.session['user_id'] = user.id
            user.backend = 'django.contrib.auth.backends.ModelBackend' 
            login(request, user)  
            if user.role.role_name == 'buyer':
                return redirect('buyer_dashboard')
            elif user.role.role_name == 'seller':
                return redirect('seller_dashboard')
            elif user.role.role_name == 'admin':
                return redirect('admin_console')
            elif user.role.role_name == 'warehouse_manager':
                return redirect('warehouse_dashboard')
            elif user.role.role_name == 'support':
                return redirect('support_dashboard')
        else:
            return render(request, 'shop/login.html', {'error': 'Invalid username or password'})
    return render(request, 'shop/login.html')

def index(request):
    return render(request, 'shop/index.html')

def dashboard(request):
    user_id = request.session.get('user_id')
    if user_id:
        user = User.objects.get(id=user_id)
        return render(request, 'shop/dashboard.html', {'user': user})
    return redirect('login')

def logout_view(request):
    request.session.flush()  
    return redirect('login')

@role_required('buyer')
def buyer_dashboard(request):
    default_warehouse = Warehouse.objects.first()
    if not default_warehouse:
        default_warehouse = Warehouse.objects.create(address="Default Address", phonenumber="123456789")

    for product in Product.objects.all():
        ProductsOnWarehouse.objects.get_or_create(product=product, warehouse=default_warehouse)

    products = Product.objects.prefetch_related('warehouse_records')

   
    sort_by = request.GET.get('sort_by')
    order = request.GET.get('order', 'asc')

    if sort_by:
        sort_order = sort_by if order == 'asc' else f"-{sort_by}"
        try:
           
            if sort_by in ['name', 'cost']:
                products = products.order_by(sort_order)
            elif sort_by == 'category': 
                products = products.order_by(f"{sort_order}category__name")
        except Exception as e:
            print(f"Ошибка сортировки: {e}")

    return render(request, 'shop/buyer_dashboard.html', {'products': products})

@role_required('seller')
def seller_dashboard(request):
    products = Product.objects.all()
    return render(request, 'shop/seller_dashboard.html', {'products': products})

@role_required('seller')
def add_product(request):
    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES)  
        if form.is_valid():
            product = form.save()
            
            default_warehouse = Warehouse.objects.first()  
            if not default_warehouse:
                default_warehouse = Warehouse.objects.create(address="Default Address", phonenumber="123456789")
            ProductsOnWarehouse.objects.create(product=product, warehouse=default_warehouse, quantity=0)
            
            return redirect('seller_dashboard')
    else:
        form = ProductForm()
    return render(request, 'shop/add_product.html', {'form': form})


@role_required('seller')
def edit_product(request, product_id):
    product = Product.objects.get(id=product_id)
    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES, instance=product)  
        if form.is_valid():
            form.save()
            return redirect('seller_dashboard')
    else:
        form = ProductForm(instance=product)
    return render(request, 'shop/edit_product.html', {'form': form})

@role_required('seller')
def delete_product(request, product_id):
    product = Product.objects.get(id=product_id)
    product.delete()
    return redirect('seller_dashboard')

@role_required('admin')
def admin_console(request):
    context = {
        'users': User.objects.all(),
        'products': Product.objects.all(),
        'categories': Category.objects.all(),
        'orders': Order.objects.all(),
        'support_messages': SupportMessage.objects.all(),
        'roles': Role.objects.all(),
        'payment_methods': PaymentMethod.objects.all(),
        'warehouses': Warehouse.objects.all(),
        'reviews': Review.objects.all(),
        'products_on_warehouse': ProductsOnWarehouse.objects.all(),
        'what_in_order': WhatInOrder.objects.all(),
    }
    return render(request, 'shop/admin_console.html', context)






def add_to_cart(request, product_id):
    user = get_object_or_404(User, id=request.session.get('user_id'))
    product = get_object_or_404(Product, id=product_id)

    stock = ProductsOnWarehouse.objects.filter(product=product).first()
    if not stock or stock.quantity < 1:
        return JsonResponse({'success': False, 'message': 'Товар закончился.'})

    cart_item, created = Cart.objects.get_or_create(user=user, product=product)
    if not created:
        cart_item.quantity += 1
    else:
        cart_item.quantity = 1

    cart_item.save()

    stock.quantity -= 1
    stock.save()

    return JsonResponse({'success': True, 'message': f'{product.name} добавлен в корзину.'})

def cart(request):
    user = get_object_or_404(User, id=request.session.get('user_id'))
    cart_items = Cart.objects.filter(user=user)

    for item in cart_items:
        item.total_cost = item.product.cost * item.quantity

    total_cost = sum(item.total_cost for item in cart_items)

    return render(request, 'shop/cart.html', {
        'cart_items': cart_items,
        'total_cost': total_cost
    })

def remove_from_cart(request, cart_id):
    cart_item = Cart.objects.get(id=cart_id)
    cart_item.delete()
    return redirect('cart')


def profile(request):
    user = User.objects.get(id=request.session['user_id'])
    orders = Order.objects.filter(user=user)

    return render(request, 'shop/profile.html', {
        'user': user,
        'orders': orders
    })

@login_required
def staff_profile(request):
    user = request.user
    role_name = user.role.role_name 

    context = {
        'user_name': user.user_name,
        'role_name': role_name
    }

   
    if role_name == 'admin':
        context['tasks'] = "Управляйте пользователями, просматривайте отчёты и заказы."
        context['permissions'] = "Вы можете добавлять, редактировать и удалять любые данные."
    elif role_name == 'seller':
        context['tasks'] = "Обрабатывайте заказы и консультируйте клиентов."
        context['permissions'] = "Вы можете просматривать и обновлять заказы."
    elif role_name == 'warehouse_manager':
        context['tasks'] = "Обрабатывайте количество товара на складе."
        context['permissions'] = "Вы можете просматривать и обновлять количество товара на складе."
    elif role_name == 'support':
        context['tasks'] = "Обрабатывайте вопросы клиентов."
        context['permissions'] = "Вы можете просматривать и отвечать на вопросы клиентов."

    else:
        context['tasks'] = "Роль не определена. Обратитесь к администратору."
        context['permissions'] = "Доступ ограничен."

    return render(request, 'shop/staff_profile.html', context) 

@role_required('buyer')
def checkout(request):
    import requests

    user = get_object_or_404(User, id=request.session.get('user_id')) 
    cart_items = Cart.objects.filter(user=user)

    if not cart_items.exists():
        return redirect('cart')  

    total_cost = sum(item.product.cost * item.quantity for item in cart_items)
    payment_methods = PaymentMethod.objects.all()

    if request.method == 'POST':
        address = request.POST.get('address', 'Default Address')
        selected_payment_id = request.POST.get('payment_method')
        selected_payment = get_object_or_404(PaymentMethod, id=selected_payment_id)

   
        order = Order.objects.create(
            user=user,
            date_order=now().date(),
            status=OrderStatus.objects.get_or_create(status="Pending")[0],
            address_order=address,
            finish_cost=total_cost,
            payment=selected_payment
        )

       
        for item in cart_items:
            WhatInOrder.objects.create(
                product=item.product,
                order=order,
                quantity=item.quantity,
                cost=item.product.cost * item.quantity
            )


        cart_items.delete()

  
        formspree_data = {
            'name': user.user_name,
            'email': user.email,
            'total_cost': f'{total_cost} ₽',
            'address': address,
            'payment_method': selected_payment.payment_method,
        }
        try:
            formspree_endpoint = 'https://formspree.io/f/xeoqzwok'
            response = requests.post(formspree_endpoint, data=formspree_data)
            if response.status_code == 200:
                logger.info("Data sent to Formspree successfully.")
            else:
                logger.warning(f"Formspree response: {response.status_code} - {response.text}")
        except Exception as e:
            logger.error(f"Error sending data to Formspree: {e}")

        return redirect('profile')

    return render(request, 'shop/checkout.html', {
        'user': user,
        'cart_items': cart_items,
        'total_cost': total_cost,
        'payment_methods': payment_methods,
    })






def update_cart_quantity(request, cart_id):
    if request.method == 'POST':
        cart_item = get_object_or_404(Cart, id=cart_id)
        new_quantity = int(request.POST.get('quantity', 1))

        if new_quantity < 1:
            logger.warning(f"Invalid quantity: {new_quantity}")
            return JsonResponse({'success': False, 'error': 'Quantity must be at least 1.'})

        cart_item.quantity = new_quantity
        cart_item.save()
        logger.info(f"Cart item {cart_item.product.name} updated to quantity {new_quantity}.")

       
        total_cost = sum(item.product.cost * item.quantity for item in Cart.objects.filter(user=cart_item.user))
        logger.info(f"Total cart cost updated: {total_cost} ₽.")

        return JsonResponse({'success': True, 'total_cost': total_cost})

    return JsonResponse({'success': False, 'error': 'Invalid request method.'})



@role_required('seller')
def seller_orders(request):
    orders = Order.objects.all()
    statuses = OrderStatus.objects.all()

    if request.method == 'POST':
        order_id = request.POST.get('order_id')
        new_status = request.POST.get('status')
        order = get_object_or_404(Order, id=order_id)
        order.status = OrderStatus.objects.get(status=new_status)
        order.save()

    return render(request, 'shop/seller_orders.html', {'orders': orders, 'statuses': statuses})

@role_required('seller')
def update_order_status(request, order_id):
    if request.method == 'POST':
        order = get_object_or_404(Order, id=order_id)
        new_status = request.POST.get('status')
        status_obj, created = OrderStatus.objects.get_or_create(status=new_status)
        order.status = status_obj
        order.save()
        return redirect('seller_orders')



@role_required('warehouse_manager')
def warehouse_dashboard(request):
    default_warehouse = Warehouse.objects.first()
    if not default_warehouse:
        default_warehouse = Warehouse.objects.create(address="Default Address", phonenumber="123456789")

    for product in Product.objects.all():
        ProductsOnWarehouse.objects.get_or_create(product=product, warehouse=default_warehouse)

    products_on_warehouse = ProductsOnWarehouse.objects.select_related('product', 'warehouse').all()

    return render(request, 'shop/warehouse_dashboard.html', {
        'products_on_warehouse': products_on_warehouse
    })




@role_required('warehouse_manager')
def update_stock(request, product_id):
    if request.method == 'POST':
        product_on_warehouse = get_object_or_404(ProductsOnWarehouse, product__id=product_id)
        action = request.POST.get('action')  
        quantity = int(request.POST.get('quantity', 0))

        if action == 'add':
            product_on_warehouse.quantity = F('quantity') + quantity
        elif action == 'remove':
            product_on_warehouse.quantity = F('quantity') - quantity

        product_on_warehouse.save()
        return redirect('warehouse_dashboard')

    return JsonResponse({'success': False, 'error': 'Invalid request method'})

@role_required('warehouse_manager')
def warehouse_manage_orders(request):
    orders = Order.objects.all()
    statuses = OrderStatus.objects.all()

    return render(request, 'shop/warehouse_manage_orders.html', {
        'orders': orders,
        'statuses': statuses
    })


def support_request(request):
    if request.method == 'POST':
        form = SupportMessageForm(request.POST)
        if form.is_valid():
            support_message = form.save(commit=False)
            support_message.user = User.objects.get(id=request.session.get('user_id')) 
            support_message.save()
            return redirect('support_success')
    else:
        form = SupportMessageForm()

    user = User.objects.get(id=request.session.get('user_id'))
    messages = SupportMessage.objects.filter(user=user).order_by('-created_at')
    return render(request, 'shop/support_request.html', {
        'form': form,
        'messages': messages
    })


@role_required('support')
def support_dashboard(request):
    messages = SupportMessage.objects.filter(is_resolved=False).order_by('-created_at')
    return render(request, 'shop/support_dashboard.html', {'messages': messages})

@role_required('support')
def support_reply(request, message_id):
    message = get_object_or_404(SupportMessage, id=message_id)
    if request.method == 'POST':
        response = request.POST.get('response')
        if response:
            message.response = response
            message.is_resolved = True
            message.save()
            return redirect('support_dashboard')
    return render(request, 'shop/support_reply.html', {'message': message})

@role_required('admin')
def view_database(request):
    users = User.objects.all()
    products = Product.objects.all()
    categories = Category.objects.all()
    orders = Order.objects.all()
    support_messages = SupportMessage.objects.all()

    return render(request, 'shop/view_database.html', {
        'users': users,
        'products': products,
        'categories': categories,
        'orders': orders,
        'support_messages': support_messages
    })


@role_required('admin')
def manage_table(request, table_name):
    model = MODEL_MAP.get(table_name)
    if not model:
        return HttpResponse("Invalid table name.", status=404)

   
    fields = [field.name for field in model._meta.fields]
    objects = model.objects.all()

  
    search_query = request.GET.get('search', '').strip()
    if search_query:
        
        search_filters = Q()
        for field in model._meta.fields:
            if isinstance(field, models.CharField) or isinstance(field, models.TextField):
                search_filters |= Q(**{f"{field.name}__icontains": search_query})
            elif isinstance(field, models.ForeignKey):
               
                related_field_name = f"{field.name}__{field.related_model._meta.pk.name}"
                search_filters |= Q(**{f"{related_field_name}__icontains": search_query})
        objects = objects.filter(search_filters)

    # Фильтрация
    filter_field = request.GET.get('filter_field')
    filter_value = request.GET.get('filter_value')
    if filter_field and filter_value:
        try:
            objects = objects.filter(**{filter_field: filter_value})
        except FieldError:
            pass  

    return render(request, 'shop/manage_table.html', {
        'table_name': table_name,
        'objects': objects,
        'fields': fields,
        'search_query': search_query,
        'filter_field': filter_field,
        'filter_value': filter_value,
    })




@role_required('admin')
def add_object(request, table_name):
    model = MODEL_MAP.get(table_name)
    if not model:
        return JsonResponse({'error': 'Invalid table name'}, status=400)

    form_class = modelform_factory(model, fields="__all__")
    form = form_class()

    if request.method == 'POST':
        form = form_class(request.POST)
        if form.is_valid():
            form.save()
            return redirect('manage_table', table_name=table_name)

    return render(request, 'shop/add_object.html', {'form': form, 'table_name': table_name})

@role_required('admin')
def edit_object(request, table_name, object_id):
    model = MODEL_MAP.get(table_name)
    if not model:
        return JsonResponse({'error': 'Invalid table name'}, status=400)

    obj = get_object_or_404(model, pk=object_id)
    form_class = modelform_factory(model, fields="__all__")
    form = form_class(instance=obj)

    if request.method == 'POST':
        form = form_class(request.POST, instance=obj)
        if form.is_valid():
            form.save()
            return redirect('manage_table', table_name=table_name)

    return render(request, 'shop/edit_object.html', {'form': form, 'table_name': table_name})

@role_required('admin')
def delete_object(request, table_name, object_id):
    table_mapping = {
        'users': User,
        'products': Product,
        'categories': Category,
        'orders': Order,
        'roles': Role,
        'payment_methods': PaymentMethod,
        'warehouses': Warehouse,
        'reviews': Review,
        'products_on_warehouse': ProductsOnWarehouse,
        'support_messages': SupportMessage,
        'what_in_order': WhatInOrder,
    }

    model = table_mapping.get(table_name)
    if not model:
        return HttpResponse("Invalid table name.", status=404)

    obj = get_object_or_404(model, id=object_id)
    obj.delete()
    return redirect('manage_table', table_name=table_name)

@login_required
def add_review(request, product_id):
    product = get_object_or_404(Product, id=product_id)
    if request.method == 'POST':
        form = ReviewForm(request.POST)
        if form.is_valid():
            review = form.save(commit=False)
            review.user = request.user
            review.product = product
            review.save()
            return redirect('product_reviews', product_id=product.id)
    else:
        form = ReviewForm()
    return render(request, 'shop/cr_review.html', {'form': form, 'product': product})

def product_reviews(request, product_id):
    product = get_object_or_404(Product, id=product_id)
    reviews = product.review_set.all()  
    return render(request, 'shop/product_reviews.html', {
        'product': product,
        'reviews': reviews,
        'average_rating': product.average_rating,
    })


def generate_qr_code(order):
    """
    Генерация QR-кода для заказа.
    """
    qr_data = f"Order ID: {order.id}\n" \
              f"User: {order.user.user_name}\n" \
              f"Address: {order.address_order}\n" \
              f"Total Cost: {order.finish_cost} ₽\n" \
              f"Date: {order.date_order}\n" \
              f"Status: {order.status.status}"
    
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(qr_data)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buffer = BytesIO()
    img.save(buffer, format="PNG")
    buffer.seek(0)
    return HttpResponse(buffer, content_type="image/png")

@login_required
def profile(request):
    user = request.user
    orders = Order.objects.filter(user=user)
    return render(request, 'shop/profile.html', {
        'user': user,
        'orders': orders
    })

def generate_order_qr(request, order_id):
    order = get_object_or_404(Order, id=order_id, user=request.user)
    return generate_qr_code(order)


@role_required('admin')
def sales_chart(request):
    import matplotlib
    matplotlib.use('Agg')  
    from django.db.models import Sum
    import matplotlib.pyplot as plt
    from io import BytesIO
    import base64

    today = now().date()
    sales_data = (
        WhatInOrder.objects.filter(order__date_order=today)
        .values('product__name')  
        .annotate(total_quantity=Sum('quantity')) 
        .order_by('-total_quantity')[:5]  
    )

    product_names = [item['product__name'] for item in sales_data]
    quantities = [item['total_quantity'] for item in sales_data]

    plt.figure(figsize=(10, 6))
    plt.bar(product_names, quantities, color='skyblue')
    plt.title("Топ-5 товаров по продажам за день")
    plt.xlabel("Товары")
    plt.ylabel("Количество продаж")
    plt.xticks(rotation=45)
    plt.tight_layout()

    buffer = BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    image_png = buffer.getvalue()
    buffer.close()

    chart = base64.b64encode(image_png).decode('utf-8')

    return render(request, 'shop/sales_chart.html', {'chart': chart})


def export_sales_report(request):
    document = Document()
    document.add_heading('Отчет о продажах', level=1)

    orders = Order.objects.all()
    table = document.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID Заказа'
    hdr_cells[1].text = 'Дата заказа'
    hdr_cells[2].text = 'Пользователь'
    hdr_cells[3].text = 'Сумма заказа'
    hdr_cells[4].text = 'Метод оплаты'

    for order in orders:
        row_cells = table.add_row().cells
        row_cells[0].text = str(order.id)
        row_cells[1].text = order.date_order.strftime('%d.%m.%Y')

        user_display = getattr(order.user, 'username', None) or getattr(order.user, 'email', 'Неизвестный пользователь')
        row_cells[2].text = user_display

        row_cells[3].text = f'{order.finish_cost} ₽'

        row_cells[4].text = getattr(order.payment, 'method_name', 'Неизвестный метод')

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=sales_report.docx'
    document.save(response)
    return response




import csv
from django.shortcuts import render, redirect
from django.http import HttpResponse
from .models import User, Order, PaymentMethod, OrderStatus

def import_orders(request):
    if request.method == "POST":
        csv_file = request.FILES.get("csv_file")

        if not csv_file or not csv_file.name.endswith('.csv'):
            return HttpResponse("Загрузите корректный CSV-файл.")

        try:
            try:
                decoded_file = csv_file.read().decode('utf-8')
            except UnicodeDecodeError:
                decoded_file = csv_file.read().decode('windows-1251') 

            reader = csv.DictReader(decoded_file.splitlines())

            for row in reader:
                user, _ = User.objects.get_or_create(username=row.get('Пользователь', 'Неизвестный пользователь'))
                payment_method, _ = PaymentMethod.objects.get_or_create(
                    method_name=row.get('Метод оплаты', 'Неизвестный метод')
                )
                status, _ = OrderStatus.objects.get_or_create(status=row.get('Статус', 'Неизвестный статус'))

                Order.objects.create(
                    user=user,
                    date_order=row.get('Дата заказа', None),
                    finish_cost=float(row.get('Сумма заказа', 0)),
                    address_order=row.get('Адрес', ''),
                    status=status,
                    payment=payment_method
                )

            return redirect('admin_console')

        except Exception as e:
            return HttpResponse(f"Ошибка при обработке файла: {str(e)}")

    return render(request, "shop/import_csv.html")


from openpyxl import Workbook
from django.http import HttpResponse
from .models import Order

import csv
from django.http import HttpResponse
from .models import Order

def export_sales_report_csv(request):
    response = HttpResponse(content_type='text/csv; charset=utf-8')
    response['Content-Disposition'] = 'attachment; filename="sales_report.csv"'

    response.write('\ufeff')  # BOM для Excel
    writer = csv.writer(response, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)

    headers = ["ID Заказа", "Дата заказа", "Пользователь", "Сумма заказа", "Метод оплаты"]
    writer.writerow(headers)

    orders = Order.objects.all()
    for order in orders:
        user_display = getattr(order.user, 'username', None) or getattr(order.user, 'email', 'Неизвестный пользователь')
        payment_method = getattr(order.payment, 'payment_method', 'Неизвестный метод')

        writer.writerow([
            order.id,
            order.date_order.strftime('%d.%m.%Y') if order.date_order else 'Дата неизвестна',
            user_display,
            f"{order.finish_cost}",
            payment_method
        ])

    return response

